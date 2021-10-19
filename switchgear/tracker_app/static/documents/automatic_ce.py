from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tracker_app.static import documents

from tracker_app.models import Switchgear, Company


def automatic_ce_gen(request, switchgear_id):
    switchgear = Switchgear.objects.get(pk=switchgear_id)
    company = Company.objects.get(pk=1)
    doc_date = switchgear.actual_shipment
    switchgear_number = switchgear.serial_no
    switchgear_name = switchgear.name

    document = Document()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=ce_{switchgear_name}.docx'
    section = document.sections[0]
    header = section.header
    paragraph_header_logo = header.paragraphs[0]
    logo_run = paragraph_header_logo.add_run()
    logo_run.add_picture(f"{company.logo.path}", width=Inches(2.5))
    paragraph_header_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph_place = document.add_paragraph(
        f"""Miejsce i data wystawienia dokumentu:                                   Warszawa, dn. {doc_date} r.""")
    paragraph_place_format = paragraph_place.paragraph_format
    paragraph_place_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph_switchgear_number = document.add_paragraph(f"""

        Deklaracja zgodności Nr. {switchgear_number}
        """)
    paragraph_switchgear_number_format = paragraph_switchgear_number.paragraph_format
    paragraph_switchgear_number_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("""

        Producent: RGT-Electro Cezary Ryll
        """)
    document.add_paragraph(
        'Dotyczy: Umowa 2020/RGT/01', style='List Bullet'
    )
    document.add_paragraph(f"""Opis produktu (rozdzielnicy): Szafa {switchgear_name}
        Zgodność z normą:""")
    document.add_paragraph(
        'PN-EN 61439-1 2011 rok (Rozdzielnice i sterownice niskonapięciowe),', style='List Number'
    )
    document.add_paragraph(
        'PN-EN 61439.2 2012 rok (Rozdzielnice i sterowanie do rozdziału energii elektrycznej PSC),', style='List Number'
    )
    document.add_paragraph(
        'PN-EN 61439.3 2011 rok (Rozdzielnice tablicowe DBOs),', style='List Number'
    )
    document.add_paragraph(
        'PN-EN 61439.4 2011 rok (Zestawy przeznaczone do instalowania na terenach budowy),', style='List Number'
    )
    document.add_paragraph(
        'PN-EN 61439.5 2011 rok (Rozdzielnice Szafowe).', style='List Number'
    )
    document.add_paragraph("""
        Niniejszym deklarujemy, że w/w wyrób jest zgodny z postanowieniami 
        PN-EN 61439 2011 rok.
        Do wykonania rozdzielnicy zostały użyte aparaty firm: Bekchoff, Eaton, Ergom, Mikrotik, Weidmueller, oraz rozdzielnica firmy Ergom IP 65. 
        Poświadczam zgodność z powyższymi normami""")
    document.add_paragraph(f"""

        \t\t\t\t\t\t\t\t\t Właściciel
        \t\t\t\t\t\t\t\t\t{company.owner}""")

    section = document.sections[0]
    footer = section.footer
    paragraph_footer_logo = footer.paragraphs[0]
    logo_run = paragraph_footer_logo.add_run()
    logo_run.add_picture(f"{company.logo.path}", width=Inches(2.5))
    paragraph_footer_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # document.add_page_break()
    document.save(response)

    return response
