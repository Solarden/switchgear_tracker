from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_date = '20.02.2020'

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Carlito'
font.size = Pt(9)
# document.line_spacing = 1
# style.no_space_between_paragraphs_of_same_style = True
# document.styles.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE

# HEADER SECTION
section = document.sections[0]
header = section.header
paragraph_header_logo = header.paragraphs[0]
logo_run = paragraph_header_logo.add_run()
logo_run.add_picture("placeholder.jpeg", width=Inches(2.5))
paragraph_header_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# DOC PLACE AND DATE
paragraph_date = document.add_paragraph(f"""Warszawa, dn. {doc_date} r.""")
paragraph_date_format = paragraph_date.paragraph_format
paragraph_date_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# BLANK SPACE
# document.add_paragraph('\n')

# CLIENT BASIC INFO
d1 = document.add_paragraph('Zleceniodawca –  placeholder', style='List Bullet')
paragraph_format = d1.paragraph_format
paragraph_format.line_spacing = 1
d2 = document.add_paragraph('Wykonawca – placeholder', style='List Bullet')
paragraph_format = d2.paragraph_format
paragraph_format.line_spacing = 1
d3 = document.add_paragraph('Dotyczy – placeholder', style='List Bullet')
paragraph_format = d3.paragraph_format
paragraph_format.line_spacing = 1

# BLANK SPACE
# document.add_paragraph('\n')

# SWITCHGEAR BASIC INFO
p = document.add_paragraph()
p.add_run("""ROK PRODUKCJI: placeholder
NR FABRYCZNY: placeholder
TYP: placeholder""").font.size = Pt(10)
# paragraph_format = p.paragraph_format
# paragraph_format.line_spacing = Pt(5)
# paragraph_format.space_after = Pt(3)

# BLANK SPACE
# document.add_paragraph('\n')

# TERMS OF GUARANTEE
paragraph_guarantee_format = document.add_paragraph('Warunki Gwarancji')
paragraph_guarantee_format = paragraph_guarantee_format.paragraph_format
paragraph_guarantee_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1
d1 = document.add_paragraph(
    'Gwarant udziela gwarancji:', style='List Number'
)
paragraph_format = d1.paragraph_format
paragraph_format.line_spacing = 1
d1a = document.add_paragraph(
    'Na w/w wyrób na okres 24 miesięcy od daty uruchomienia lecz nie dłużej niż 30 miesięcy od daty sprzedaży,',
    style='List Bullet 2'
)
paragraph_format = d1a.paragraph_format
paragraph_format.line_spacing = 1
d1b = document.add_paragraph(
    'Na aparaturę i urządzenia poddostawców wg. Kart Gwarancyjnych dostarczanych przez producentów na dany wyrób.',
    style='List Bullet 2'
)
paragraph_format = d1b.paragraph_format
paragraph_format.line_spacing = 1

# 2
document.add_paragraph(
    'W okresie gwarancji Gwarant zobowiązuje się do usunięcia wady fizycznej sprzedanego urządzenia jeżeli ta wada ujawni się w okresie związania gwarancją, wskazanego w ust. 1.',
    style='List Number'
)

# 3
document.add_paragraph(
    'Naprawa nastąpi w terminie i na warunkach uzgodnionych przez strony.', style='List Number'
)

# 4
document.add_paragraph(
    'W uzasadnionych przypadkach okres naprawy gwarancyjnej może ulec wydłużeniu. Gwarant zobowiązany jest do powiadomienia o terminie naprawy gwarancyjnej.',
    style='List Number'
)

# 5
d5 = document.add_paragraph(
    'Gwarant zwolniony jest z tytułu wad fizycznych, jeżeli powstały one na wskutek:', style='List Number'
)
# paragraph_format = d5.paragraph_format
# paragraph_format.line_spacing = 1
d5a = document.add_paragraph(
    'Uszkodzeń powstałych wskutek używania urządzeń niezgodnie z przeznaczeniem,',
    style='List Bullet 2'
)
paragraph_format = d5a.paragraph_format
paragraph_format.line_spacing = 1
d5bb = document.add_paragraph(
    'Uszkodzeń powstałych po wydaniu urządzenia użytkownikowi z przyczyn niezależnych od Gwaranta w szczególności od zdarzeń losowych i działania sił wyższych, lub działań osób niezależnych od Gwaranta jeżeli przyczyny te spowodowały trwałe zmiany jakościowe gwarantowanego wyrobu,',
    style='List Bullet 2'
)
paragraph_format = d5bb.paragraph_format
paragraph_format.line_spacing = 1
d5b = document.add_paragraph(
    'Uszkodzeń powstałych na skutek zmian i przeróbek wyrobu bez uzgodnień z Gwarantem,',
    style='List Bullet 2'
)
paragraph_format = d5b.paragraph_format
paragraph_format.line_spacing = 1
d5c = document.add_paragraph(
    'Uszkodzeń mechanicznych powstałych w czasie rozładunku, montażu i rozruchu urządzenia,',
    style='List Bullet 2'
)
paragraph_format = d5c.paragraph_format
paragraph_format.line_spacing = 1
d5d = document.add_paragraph(
    'Uszkodzeń po wykryciu wady  i niezgłoszonych Gwarantowi, powodującej poważniejsze wady urządzenia,',
    style='List Bullet 2'
)
paragraph_format = d5d.paragraph_format
paragraph_format.line_spacing = 1
d5e = document.add_paragraph(
    'Uszkodzeń spowodowanym używaniem urządzeń z innym niesprawnym lub uszkodzonym urządzeniem.',
    style='List Bullet 2'
)
paragraph_format = d5e.paragraph_format
paragraph_format.line_spacing = 1

# 6
document.add_paragraph(
    'Gwarant nie ponosi odpowiedzialności z tytułu gwarancji jeżeli Użytkownik nie umożliwi Gwarantowi dostępu do urządzenia z zachowaniem przepisów BHP, oraz nie zapewni odpowiedniego sprzętu (dźwig, podnośnik z koszem itp.) niezbędnego do usunięcia wady, w terminach określonych w Karcie Gwarancyjnej.',
    style='List Number'
)

# 7
document.add_paragraph(
    'Gwarant nie ponosi odpowiedzialności finansowej z tytułu przygotowania miejsca pracy, dopuszczeń i nadzorów niezbędnych do usunięcia awarii.',
    style='List Number'
)

# 8
d8 = document.add_paragraph(
    'Użytkownik traci prawo gwarancji w przypadkach:', style='List Number'
)
paragraph_format = d8.paragraph_format
paragraph_format.line_spacing = 1
d8a = document.add_paragraph(
    'Nieprzestrzegania instrukcji obsługi i przepisów eksploatacyjnych urządzeń elektro-energetycznych przy uruchamianiu, obsłudze, konserwacji i eksploatacji urządzenia,',
    style='List Bullet 2'
)
paragraph_format = d8a.paragraph_format
paragraph_format.line_spacing = 1
d8b = document.add_paragraph(
    'Samowolnego dokonywania napraw urządzenia przez osoby nieupoważnione i nieuprawnione.',
    style='List Bullet 2'
)
paragraph_format = d8b.paragraph_format
paragraph_format.line_spacing = 1

# 9
document.add_paragraph(
    'Materiały eksploatacyjne w szczególności: żarówki, bezpieczniki, diody sygnalizacyjne, podkładki izolacyjne dzielnika napięcia, itp. Nie są objęte gwarancją.',
    style='List Number'
)

# 10
document.add_paragraph(
    'Użytkownik winien zgłosić wadę urządzenia na piśmie, pocztą elektroniczną w ciągu 48 godzin od daty wydania urządzenia. W przypadku wad widocznych, niezwłocznie po ich wykryciu, w przypadku wad ukrytych nie później niż w ciągu 48 godzin od daty ich ujawnienia.',
    style='List Number'
)

# 11
document.add_paragraph(
    'Użytkownik jest zobowiązany podać w zgłoszeniu termin udostępnienia urządzenia objętego gwarancja do naprawy, oraz opis wady.',
    style='List Number'
)

# 12
document.add_paragraph(
    'Gwarant w uzasadnionych przypadkach może zarządzać odesłania urządzenia, lub wadliwej części do Gwaranta, lub na inny wskazany adres, środkiem transportu określonym przez Gwaranta.',
    style='List Number'
)

# 13
document.add_paragraph(
    'Gwarant zobowiązuje się do odesłania wolnego od wad urządzenia na swój koszt.', style='List Number'
)

# 14
document.add_paragraph(
    'W przypadku braku możliwości dostarczenia urządzenia objętego gwarancją do siedziby Gwaranta, lub na inny wskazany przez Gwaranta adres, Gwarant zobowiązuje się do: dokonania wizji lokalnych, naprawy, wymiany w miejscu zainstalowania.',
    style='List Number'
)

# 15
document.add_paragraph(
    'Gwarant zobowiązuję się podjąć czynności związane z naprawą urządzenia niezwłocznie po zgłoszeniu na piśmie awarii urządzenia.',
    style='List Number'
)

# 16
document.add_paragraph(
    'Uprawnienia z tytułu udzielonej gwarancji mogą być realizowane po przedstawieniu ważnej karty gwarancyjnej.',
    style='List Number'
)

# 17
document.add_paragraph(
    'W przypadku stwierdzenia przez Gwaranta, iż nastąpiło nieuzasadnione zgłoszenie przez Użytkownika wad urządzenia w ramach gwarancji, Użytkownik ponosi wszelkie koszty działań podjętych przez Gwaranta.',
    style='List Number'
)

# 18
document.add_paragraph(
    'Naprawa w miejscu zainstalowana będzie odbywać się przy udziale przedstawiciela Użytkownika.', style='List Number'
)

# 19
document.add_paragraph(
    'W kontaktach z Gwarantem i jego przedstawicielami, Użytkownika reprezentować może jedynie upoważniony przedstawiciel.',
    style='List Number'
)

# 20
document.add_paragraph(
    'Gwarant nie ponosi odpowiedzialności za szkody spowodowane wyłączeniem z eksploatacji urządzeń w okresie od ujawnienia usterki, lub wady do czasu jej usunięcia, oraz szkody następcze lub pośrednie, w tym za utracone korzyści, spowodowane wystąpieniem wady urządzenia.',
    style='List Number'
)

# 21
document.add_paragraph(
    'Gwarancja obowiązuje na terenie Rzeczpospolitej Polskiej.', style='List Number'
)

# 22
d22 = document.add_paragraph(
    'Awarie urządzeń należy zgłaszać w dni robocze w godzinach 6.00-16.00', style='List Number'
)
paragraph_format = d22.paragraph_format
paragraph_format.line_spacing = 1
d23 = document.add_paragraph(
    'Numer telefonu placeholder,',
    style='List Bullet 2'
)
paragraph_format = d23.paragraph_format
paragraph_format.line_spacing = 1
d24 = document.add_paragraph(
    'E-mail: placeholder.',
    style='List Bullet 2'
)
paragraph_format = d24.paragraph_format
paragraph_format.line_spacing = 1
d25 = document.add_paragraph(
    'W przypadkach niecierpiących zwłoki – poza godzinami pracy 24/h placeholder'
)
paragraph_format = d25.paragraph_format
paragraph_format.line_spacing = 1

# FOOTER SECTION
section = document.sections[0]
footer = section.footer
paragraph_footer_logo = footer.paragraphs[0]
logo_run = paragraph_footer_logo.add_run()
logo_run.add_picture("rgt_logo.jpeg", width=Inches(2.5))
paragraph_footer_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# document.add_page_break()

document.save('guarantee_demo1.docx')
